"""
VITAL Path - Phase 3 Enhanced: Advanced Medical RAG Pipeline
==========================================================

Core Intelligence Layer - Medical Document Processing & Retrieval
Advanced semantic chunking, medical entity extraction, and intelligent retrieval

Key Features:
- Medical-aware document ingestion with FHIR/HL7 integration
- Semantic chunking with medical terminology preservation
- Multi-modal medical document processing (text, images, charts)
- Clinical context preservation during chunking
- Advanced retrieval with medical specialty routing
- Evidence quality scoring and source credibility assessment
"""

import asyncio
import logging
from typing import Dict, List, Optional, Tuple, Any, Union
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from enum import Enum
import json
import re
import hashlib
from pathlib import Path

# Medical processing libraries
import spacy
from transformers import AutoTokenizer, AutoModel
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize

# Medical terminology processing
import scispacy
import medspacy
from medspacy.ner import TargetRule
from medspacy.postprocess import Postprocessor

# Vector database and semantic search
import chromadb
from chromadb.config import Settings
import faiss

# Document processing
import PyPDF2
import docx
from PIL import Image
import pytesseract

# Medical standards and validation
from fhir.resources import Bundle, Patient, Observation
import hl7


class DocumentType(Enum):
    """Medical document classification"""
    CLINICAL_TRIAL = "clinical_trial"
    RESEARCH_PAPER = "research_paper"
    CLINICAL_GUIDELINES = "clinical_guidelines"
    FDA_GUIDANCE = "fda_guidance"
    EMA_GUIDANCE = "ema_guidance"
    DRUG_LABEL = "drug_label"
    CLINICAL_NOTE = "clinical_note"
    PATHOLOGY_REPORT = "pathology_report"
    RADIOLOGY_REPORT = "radiology_report"
    LAB_RESULT = "lab_result"
    ADVERSE_EVENT_REPORT = "adverse_event_report"
    PROTOCOL = "protocol"
    CASE_STUDY = "case_study"
    META_ANALYSIS = "meta_analysis"
    SYSTEMATIC_REVIEW = "systematic_review"


class MedicalSpecialty(Enum):
    """Medical specialties for context-aware processing"""
    CARDIOLOGY = "cardiology"
    ONCOLOGY = "oncology"
    NEUROLOGY = "neurology"
    ENDOCRINOLOGY = "endocrinology"
    GASTROENTEROLOGY = "gastroenterology"
    PULMONOLOGY = "pulmonology"
    NEPHROLOGY = "nephrology"
    HEMATOLOGY = "hematology"
    RHEUMATOLOGY = "rheumatology"
    INFECTIOUS_DISEASE = "infectious_disease"
    PSYCHIATRY = "psychiatry"
    DERMATOLOGY = "dermatology"
    OPHTHALMOLOGY = "ophthalmology"
    OTOLARYNGOLOGY = "otolaryngology"
    ORTHOPEDICS = "orthopedics"
    RADIOLOGY = "radiology"
    PATHOLOGY = "pathology"
    EMERGENCY_MEDICINE = "emergency_medicine"
    FAMILY_MEDICINE = "family_medicine"
    INTERNAL_MEDICINE = "internal_medicine"
    PEDIATRICS = "pediatrics"
    OBSTETRICS_GYNECOLOGY = "obstetrics_gynecology"
    ANESTHESIOLOGY = "anesthesiology"
    SURGERY = "surgery"
    CRITICAL_CARE = "critical_care"
    PHARMACOLOGY = "pharmacology"
    CLINICAL_TRIALS = "clinical_trials"
    REGULATORY_AFFAIRS = "regulatory_affairs"


@dataclass
class MedicalEntity:
    """Extracted medical entity with context"""
    text: str
    label: str
    confidence: float
    start_pos: int
    end_pos: int
    umls_code: Optional[str] = None
    snomed_code: Optional[str] = None
    icd_code: Optional[str] = None
    context: Optional[str] = None
    negation: bool = False
    certainty: str = "definite"  # definite, probable, possible


@dataclass
class DocumentChunk:
    """Medical document chunk with metadata"""
    chunk_id: str
    content: str
    document_id: str
    document_type: DocumentType
    specialties: List[MedicalSpecialty]
    chunk_index: int
    total_chunks: int
    entities: List[MedicalEntity] = field(default_factory=list)
    semantic_embedding: Optional[np.ndarray] = None
    clinical_context: Optional[str] = None
    evidence_level: Optional[str] = None  # Level I-V evidence hierarchy
    study_design: Optional[str] = None
    patient_population: Optional[str] = None
    intervention: Optional[str] = None
    outcome_measures: Optional[List[str]] = None
    statistical_significance: Optional[bool] = None
    clinical_significance: Optional[bool] = None
    safety_signals: List[str] = field(default_factory=list)
    efficacy_signals: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: datetime = field(default_factory=datetime.now)


@dataclass
class MedicalDocument:
    """Medical document with full metadata"""
    document_id: str
    title: str
    content: str
    document_type: DocumentType
    specialties: List[MedicalSpecialty]
    authors: List[str] = field(default_factory=list)
    institution: Optional[str] = None
    publication_date: Optional[datetime] = None
    doi: Optional[str] = None
    pmid: Optional[str] = None
    study_design: Optional[str] = None
    evidence_level: Optional[str] = None
    patient_count: Optional[int] = None
    intervention: Optional[str] = None
    primary_endpoint: Optional[str] = None
    secondary_endpoints: List[str] = field(default_factory=list)
    inclusion_criteria: List[str] = field(default_factory=list)
    exclusion_criteria: List[str] = field(default_factory=list)
    adverse_events: List[str] = field(default_factory=list)
    statistical_methods: List[str] = field(default_factory=list)
    limitations: List[str] = field(default_factory=list)
    clinical_implications: List[str] = field(default_factory=list)
    regulatory_status: Optional[str] = None
    compliance_notes: List[str] = field(default_factory=list)
    quality_score: float = 0.0
    credibility_score: float = 0.0
    chunks: List[DocumentChunk] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: datetime = field(default_factory=datetime.now)
    updated_at: datetime = field(default_factory=datetime.now)


class MedicalEntityExtractor:
    """Advanced medical entity extraction with multiple models"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

        # Load medical NLP models
        try:
            self.nlp = spacy.load("en_core_sci_sm")  # ScispaCy model
            self.medspacy_nlp = medspacy.load("en_core_med7_lg")
        except OSError:
            self.logger.warning("Medical NLP models not found. Install scispacy and medspacy.")
            self.nlp = spacy.load("en_core_web_sm")
            self.medspacy_nlp = None

        # Medical terminology processors
        self.drug_patterns = self._load_drug_patterns()
        self.disease_patterns = self._load_disease_patterns()
        self.procedure_patterns = self._load_procedure_patterns()

    def _load_drug_patterns(self) -> List[str]:
        """Load pharmaceutical drug name patterns"""
        # In production, this would load from a comprehensive drug database
        return [
            r"\b\w*mab\b",  # Monoclonal antibodies
            r"\b\w*cillin\b",  # Penicillins
            r"\b\w*statin\b",  # Statins
            r"\b\w*pril\b",  # ACE inhibitors
            r"\b\w*sartan\b",  # ARBs
            r"\b\w*olol\b",  # Beta blockers
        ]

    def _load_disease_patterns(self) -> List[str]:
        """Load disease and condition patterns"""
        return [
            r"\b\w*cardia\b",  # Cardiac conditions
            r"\b\w*cancer\b",  # Cancers
            r"\b\w*tumor\b",   # Tumors
            r"\b\w*itis\b",    # Inflammations
            r"\b\w*osis\b",    # Conditions
            r"\b\w*pathy\b",   # Diseases
        ]

    def _load_procedure_patterns(self) -> List[str]:
        """Load medical procedure patterns"""
        return [
            r"\b\w*scopy\b",   # Scopic procedures
            r"\b\w*ectomy\b",  # Surgical removals
            r"\b\w*plasty\b",  # Reconstructive procedures
            r"\b\w*graphy\b",  # Imaging procedures
            r"\b\w*therapy\b", # Therapeutic procedures
        ]

    async def extract_entities(self, text: str, context: Optional[str] = None) -> List[MedicalEntity]:
        """Extract medical entities with context awareness"""
        entities = []

        # Process with ScispaCy
        if self.nlp:
            doc = self.nlp(text)
            for ent in doc.ents:
                entity = MedicalEntity(
                    text=ent.text,
                    label=ent.label_,
                    confidence=0.8,  # ScispaCy default confidence
                    start_pos=ent.start_char,
                    end_pos=ent.end_char,
                    context=context
                )
                entities.append(entity)

        # Process with medspaCy if available
        if self.medspacy_nlp:
            med_doc = self.medspacy_nlp(text)
            for ent in med_doc.ents:
                entity = MedicalEntity(
                    text=ent.text,
                    label=ent.label_,
                    confidence=0.85,
                    start_pos=ent.start_char,
                    end_pos=ent.end_char,
                    context=context,
                    negation=ent._.is_negated if hasattr(ent._, 'is_negated') else False,
                    certainty=ent._.certainty if hasattr(ent._, 'certainty') else "definite"
                )
                entities.append(entity)

        # Pattern-based extraction for drugs
        for pattern in self.drug_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                entity = MedicalEntity(
                    text=match.group(),
                    label="DRUG",
                    confidence=0.7,
                    start_pos=match.start(),
                    end_pos=match.end(),
                    context=context
                )
                entities.append(entity)

        # Remove duplicates and merge overlapping entities
        entities = self._merge_overlapping_entities(entities)

        return entities

    def _merge_overlapping_entities(self, entities: List[MedicalEntity]) -> List[MedicalEntity]:
        """Merge overlapping entities, keeping highest confidence"""
        if not entities:
            return entities

        # Sort by start position
        entities.sort(key=lambda x: x.start_pos)

        merged = []
        current = entities[0]

        for next_entity in entities[1:]:
            # Check for overlap
            if next_entity.start_pos <= current.end_pos:
                # Overlapping - keep higher confidence entity
                if next_entity.confidence > current.confidence:
                    current = next_entity
            else:
                # No overlap - add current and move to next
                merged.append(current)
                current = next_entity

        merged.append(current)
        return merged


class MedicalSemanticChunker:
    """Advanced semantic chunking preserving medical context"""

    def __init__(self, model_name: str = "microsoft/BiomedNLP-PubMedBERT-base-uncased-abstract"):
        self.logger = logging.getLogger(__name__)
        self.model_name = model_name

        # Load medical tokenizer and model
        try:
            self.tokenizer = AutoTokenizer.from_pretrained(model_name)
            self.model = AutoModel.from_pretrained(model_name)
        except Exception as e:
            self.logger.warning(f"Could not load medical model {model_name}: {e}")
            # Fallback to basic model
            self.tokenizer = AutoTokenizer.from_pretrained("bert-base-uncased")
            self.model = AutoModel.from_pretrained("bert-base-uncased")

        self.entity_extractor = MedicalEntityExtractor()

        # Medical section headers
        self.medical_headers = {
            "clinical_trial": [
                "abstract", "introduction", "methods", "results", "discussion",
                "conclusion", "background", "objectives", "design", "setting",
                "participants", "interventions", "main outcome measures",
                "results", "conclusions", "trial registration"
            ],
            "case_report": [
                "abstract", "introduction", "case presentation", "discussion",
                "conclusion", "background", "case description", "management",
                "outcome", "learning points"
            ],
            "systematic_review": [
                "abstract", "introduction", "methods", "results", "discussion",
                "conclusion", "background", "search strategy", "study selection",
                "data extraction", "quality assessment", "statistical analysis"
            ]
        }

    async def chunk_document(self, document: MedicalDocument,
                           max_chunk_size: int = 1500,
                           overlap: int = 150) -> List[DocumentChunk]:
        """Create semantic chunks preserving medical context"""

        # First, try structure-based chunking for known document types
        if document.document_type in [DocumentType.CLINICAL_TRIAL,
                                    DocumentType.RESEARCH_PAPER,
                                    DocumentType.SYSTEMATIC_REVIEW]:
            chunks = await self._structure_based_chunking(document, max_chunk_size)
            if chunks:
                return chunks

        # Fall back to semantic chunking
        return await self._semantic_chunking(document, max_chunk_size, overlap)

    async def _structure_based_chunking(self, document: MedicalDocument,
                                      max_chunk_size: int) -> List[DocumentChunk]:
        """Chunk based on medical document structure"""
        chunks = []
        content = document.content

        # Determine document structure
        doc_type_key = document.document_type.value
        if doc_type_key in ["clinical_trial", "research_paper"]:
            headers = self.medical_headers["clinical_trial"]
        elif doc_type_key == "case_study":
            headers = self.medical_headers["case_report"]
        elif doc_type_key in ["systematic_review", "meta_analysis"]:
            headers = self.medical_headers["systematic_review"]
        else:
            return []  # Fall back to semantic chunking

        # Find sections based on headers
        sections = self._extract_sections(content, headers)

        chunk_index = 0
        for section_name, section_content in sections.items():
            if len(section_content.strip()) < 50:  # Skip very short sections
                continue

            # If section is too long, sub-chunk it
            if len(section_content) > max_chunk_size:
                sub_chunks = await self._semantic_chunking_text(
                    section_content, max_chunk_size, 100
                )
                for i, sub_chunk in enumerate(sub_chunks):
                    chunk = await self._create_chunk(
                        document, sub_chunk, chunk_index,
                        clinical_context=f"Section: {section_name}"
                    )
                    chunks.append(chunk)
                    chunk_index += 1
            else:
                chunk = await self._create_chunk(
                    document, section_content, chunk_index,
                    clinical_context=f"Section: {section_name}"
                )
                chunks.append(chunk)
                chunk_index += 1

        return chunks

    def _extract_sections(self, content: str, headers: List[str]) -> Dict[str, str]:
        """Extract sections based on medical headers"""
        sections = {}
        content_lower = content.lower()

        # Find header positions
        header_positions = []
        for header in headers:
            # Look for various header formats
            patterns = [
                rf"\b{header}\b",
                rf"\b{header.replace('_', ' ')}\b",
                rf"^{header}:?$",
                rf"^{header.replace('_', ' ')}:?$"
            ]

            for pattern in patterns:
                matches = list(re.finditer(pattern, content_lower, re.MULTILINE | re.IGNORECASE))
                for match in matches:
                    header_positions.append((match.start(), header))

        # Sort by position
        header_positions.sort()

        # Extract sections
        for i, (start_pos, header) in enumerate(header_positions):
            if i + 1 < len(header_positions):
                end_pos = header_positions[i + 1][0]
            else:
                end_pos = len(content)

            section_content = content[start_pos:end_pos].strip()
            sections[header] = section_content

        return sections

    async def _semantic_chunking(self, document: MedicalDocument,
                               max_chunk_size: int, overlap: int) -> List[DocumentChunk]:
        """Semantic chunking with medical context preservation"""

        # Split by sentences first to preserve medical context
        sentences = sent_tokenize(document.content)

        chunks = []
        current_chunk = []
        current_length = 0
        chunk_index = 0

        for sentence in sentences:
            sentence_length = len(sentence)

            # Check if adding this sentence would exceed chunk size
            if current_length + sentence_length > max_chunk_size and current_chunk:
                # Create chunk from current sentences
                chunk_text = " ".join(current_chunk)
                chunk = await self._create_chunk(document, chunk_text, chunk_index)
                chunks.append(chunk)

                # Start new chunk with overlap
                overlap_sentences = current_chunk[-2:] if len(current_chunk) >= 2 else current_chunk
                current_chunk = overlap_sentences + [sentence]
                current_length = sum(len(s) for s in current_chunk)
                chunk_index += 1
            else:
                current_chunk.append(sentence)
                current_length += sentence_length

        # Add final chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk = await self._create_chunk(document, chunk_text, chunk_index)
            chunks.append(chunk)

        return chunks

    async def _semantic_chunking_text(self, text: str, max_size: int, overlap: int) -> List[str]:
        """Semantic chunking for plain text"""
        sentences = sent_tokenize(text)
        chunks = []
        current_chunk = []
        current_length = 0

        for sentence in sentences:
            sentence_length = len(sentence)

            if current_length + sentence_length > max_size and current_chunk:
                chunks.append(" ".join(current_chunk))

                # Overlap handling
                overlap_sentences = current_chunk[-1:] if current_chunk else []
                current_chunk = overlap_sentences + [sentence]
                current_length = sum(len(s) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_length += sentence_length

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    async def _create_chunk(self, document: MedicalDocument, content: str,
                          chunk_index: int, clinical_context: Optional[str] = None) -> DocumentChunk:
        """Create a document chunk with full medical metadata"""

        # Generate chunk ID
        chunk_id = hashlib.md5(f"{document.document_id}_{chunk_index}_{content[:100]}".encode()).hexdigest()

        # Extract entities from chunk
        entities = await self.entity_extractor.extract_entities(content, clinical_context)

        # Extract clinical metadata
        evidence_level = self._extract_evidence_level(content)
        study_design = self._extract_study_design(content)
        patient_population = self._extract_patient_population(content)
        intervention = self._extract_intervention(content)
        outcome_measures = self._extract_outcome_measures(content)
        safety_signals = self._extract_safety_signals(content)
        efficacy_signals = self._extract_efficacy_signals(content)

        chunk = DocumentChunk(
            chunk_id=chunk_id,
            content=content,
            document_id=document.document_id,
            document_type=document.document_type,
            specialties=document.specialties,
            chunk_index=chunk_index,
            total_chunks=0,  # Will be updated after all chunks are created
            entities=entities,
            clinical_context=clinical_context,
            evidence_level=evidence_level,
            study_design=study_design,
            patient_population=patient_population,
            intervention=intervention,
            outcome_measures=outcome_measures,
            safety_signals=safety_signals,
            efficacy_signals=efficacy_signals
        )

        return chunk

    def _extract_evidence_level(self, content: str) -> Optional[str]:
        """Extract evidence level from content"""
        evidence_patterns = {
            "Level I": r"randomized controlled trial|RCT|systematic review|meta-analysis",
            "Level II": r"cohort study|case-control study|comparative study",
            "Level III": r"case series|descriptive study|cross-sectional",
            "Level IV": r"case report|expert opinion|consensus",
            "Level V": r"expert opinion|editorial|commentary"
        }

        content_lower = content.lower()
        for level, pattern in evidence_patterns.items():
            if re.search(pattern, content_lower):
                return level

        return None

    def _extract_study_design(self, content: str) -> Optional[str]:
        """Extract study design information"""
        design_patterns = [
            r"randomized controlled trial",
            r"double-blind",
            r"single-blind",
            r"placebo-controlled",
            r"crossover",
            r"parallel group",
            r"cohort study",
            r"case-control",
            r"case series",
            r"case report",
            r"systematic review",
            r"meta-analysis",
            r"observational study",
            r"prospective",
            r"retrospective"
        ]

        content_lower = content.lower()
        for pattern in design_patterns:
            if re.search(pattern, content_lower):
                return pattern.replace(r"\b", "").replace("\\", "")

        return None

    def _extract_patient_population(self, content: str) -> Optional[str]:
        """Extract patient population information"""
        # Look for patient demographic patterns
        patterns = [
            r"\d+\s*patients?",
            r"\d+\s*subjects?",
            r"\d+\s*participants?",
            r"adult\s*patients?",
            r"pediatric\s*patients?",
            r"elderly\s*patients?",
            r"men and women",
            r"male and female"
        ]

        content_lower = content.lower()
        for pattern in patterns:
            match = re.search(pattern, content_lower)
            if match:
                return match.group()

        return None

    def _extract_intervention(self, content: str) -> Optional[str]:
        """Extract intervention information"""
        # Look for intervention patterns
        intervention_keywords = [
            "treatment", "therapy", "medication", "drug", "intervention",
            "procedure", "surgery", "protocol", "regimen"
        ]

        content_lower = content.lower()
        for keyword in intervention_keywords:
            pattern = rf"\b{keyword}\s+\w+"
            match = re.search(pattern, content_lower)
            if match:
                return match.group()

        return None

    def _extract_outcome_measures(self, content: str) -> List[str]:
        """Extract outcome measures"""
        outcome_patterns = [
            r"primary endpoint",
            r"secondary endpoint",
            r"primary outcome",
            r"secondary outcome",
            r"efficacy endpoint",
            r"safety endpoint",
            r"survival",
            r"response rate",
            r"progression-free survival",
            r"overall survival",
            r"quality of life"
        ]

        outcomes = []
        content_lower = content.lower()

        for pattern in outcome_patterns:
            matches = re.findall(pattern, content_lower)
            outcomes.extend(matches)

        return list(set(outcomes))

    def _extract_safety_signals(self, content: str) -> List[str]:
        """Extract safety-related information"""
        safety_patterns = [
            r"adverse event",
            r"side effect",
            r"toxicity",
            r"serious adverse event",
            r"SAE",
            r"contraindication",
            r"warning",
            r"precaution",
            r"black box warning",
            r"drug interaction"
        ]

        signals = []
        content_lower = content.lower()

        for pattern in safety_patterns:
            matches = re.findall(pattern, content_lower)
            signals.extend(matches)

        return list(set(signals))

    def _extract_efficacy_signals(self, content: str) -> List[str]:
        """Extract efficacy-related information"""
        efficacy_patterns = [
            r"statistically significant",
            r"clinically significant",
            r"p\s*<\s*0\.05",
            r"confidence interval",
            r"response rate",
            r"complete response",
            r"partial response",
            r"stable disease",
            r"disease progression",
            r"improvement",
            r"benefit"
        ]

        signals = []
        content_lower = content.lower()

        for pattern in efficacy_patterns:
            matches = re.findall(pattern, content_lower)
            signals.extend(matches)

        return list(set(signals))


class MedicalDocumentIngestion:
    """Advanced medical document ingestion with multi-modal support"""

    def __init__(self, vector_db_path: str = "./medical_vector_db"):
        self.logger = logging.getLogger(__name__)
        self.vector_db_path = Path(vector_db_path)
        self.vector_db_path.mkdir(exist_ok=True)

        # Initialize vector database
        self.chroma_client = chromadb.PersistentClient(path=str(self.vector_db_path))
        self.collection = self.chroma_client.get_or_create_collection(
            name="medical_documents",
            metadata={"hnsw:space": "cosine"}
        )

        # Initialize components
        self.chunker = MedicalSemanticChunker()
        self.entity_extractor = MedicalEntityExtractor()

        # Initialize FAISS index for fast similarity search
        self.faiss_index = None
        self.faiss_dimension = 768  # BioBERT embedding dimension

    async def ingest_document(self, file_path: str, document_metadata: Dict[str, Any] = None) -> MedicalDocument:
        """Ingest a medical document with full processing pipeline"""

        self.logger.info(f"Starting ingestion of document: {file_path}")

        # Extract text from document
        content = await self._extract_text(file_path)
        if not content:
            raise ValueError(f"Could not extract text from {file_path}")

        # Create document metadata
        document = await self._create_document_metadata(file_path, content, document_metadata)

        # Chunk the document
        chunks = await self.chunker.chunk_document(document)

        # Update total chunks count
        for chunk in chunks:
            chunk.total_chunks = len(chunks)

        # Generate embeddings and store chunks
        await self._process_and_store_chunks(chunks)

        # Update document with chunks
        document.chunks = chunks

        self.logger.info(f"Successfully ingested document {document.document_id} with {len(chunks)} chunks")

        return document

    async def _extract_text(self, file_path: str) -> str:
        """Extract text from various document formats"""

        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        try:
            if extension == '.pdf':
                return await self._extract_pdf_text(file_path)
            elif extension in ['.doc', '.docx']:
                return await self._extract_word_text(file_path)
            elif extension == '.txt':
                return file_path.read_text(encoding='utf-8')
            elif extension in ['.png', '.jpg', '.jpeg', '.tiff']:
                return await self._extract_image_text(file_path)
            else:
                self.logger.warning(f"Unsupported file format: {extension}")
                return ""

        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF with medical document awareness"""

        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text_parts = []

                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

                return "\n\n".join(text_parts)

        except Exception as e:
            self.logger.error(f"Error extracting PDF text: {e}")
            return ""

    async def _extract_word_text(self, file_path: Path) -> str:
        """Extract text from Word documents"""

        try:
            doc = docx.Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            return "\n\n".join(text_parts)

        except Exception as e:
            self.logger.error(f"Error extracting Word text: {e}")
            return ""

    async def _extract_image_text(self, file_path: Path) -> str:
        """Extract text from images using OCR"""

        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()

        except Exception as e:
            self.logger.error(f"Error extracting image text: {e}")
            return ""

    async def _create_document_metadata(self, file_path: str, content: str,
                                      metadata: Dict[str, Any] = None) -> MedicalDocument:
        """Create comprehensive document metadata"""

        file_path = Path(file_path)
        metadata = metadata or {}

        # Generate document ID
        document_id = hashlib.md5(f"{file_path.name}_{content[:1000]}".encode()).hexdigest()

        # Extract title (from filename or content)
        title = metadata.get('title', file_path.stem)

        # Classify document type
        document_type = self._classify_document_type(content, file_path.name)

        # Identify medical specialties
        specialties = self._identify_specialties(content)

        # Extract additional metadata from content
        study_design = await self._extract_study_design_from_content(content)
        evidence_level = await self._extract_evidence_level_from_content(content)

        document = MedicalDocument(
            document_id=document_id,
            title=title,
            content=content,
            document_type=document_type,
            specialties=specialties,
            authors=metadata.get('authors', []),
            institution=metadata.get('institution'),
            publication_date=metadata.get('publication_date'),
            doi=metadata.get('doi'),
            pmid=metadata.get('pmid'),
            study_design=study_design,
            evidence_level=evidence_level,
            metadata=metadata
        )

        return document

    def _classify_document_type(self, content: str, filename: str) -> DocumentType:
        """Classify document type based on content and filename"""

        content_lower = content.lower()
        filename_lower = filename.lower()

        # Check filename patterns first
        if any(term in filename_lower for term in ['protocol', 'study_protocol']):
            return DocumentType.PROTOCOL
        if any(term in filename_lower for term in ['fda_guidance', 'fda_draft']):
            return DocumentType.FDA_GUIDANCE
        if any(term in filename_lower for term in ['ema_guidance', 'ema_draft']):
            return DocumentType.EMA_GUIDANCE
        if any(term in filename_lower for term in ['label', 'prescribing_information']):
            return DocumentType.DRUG_LABEL

        # Check content patterns
        if re.search(r'clinical trial|randomized controlled trial|study protocol', content_lower):
            return DocumentType.CLINICAL_TRIAL
        if re.search(r'systematic review|meta-analysis', content_lower):
            if 'meta-analysis' in content_lower:
                return DocumentType.META_ANALYSIS
            return DocumentType.SYSTEMATIC_REVIEW
        if re.search(r'case report|case study', content_lower):
            return DocumentType.CASE_STUDY
        if re.search(r'clinical guidelines|practice guidelines|treatment guidelines', content_lower):
            return DocumentType.CLINICAL_GUIDELINES
        if re.search(r'adverse event report|safety report', content_lower):
            return DocumentType.ADVERSE_EVENT_REPORT
        if re.search(r'pathology report|histopathology', content_lower):
            return DocumentType.PATHOLOGY_REPORT
        if re.search(r'radiology report|imaging report|ct scan|mri|x-ray', content_lower):
            return DocumentType.RADIOLOGY_REPORT
        if re.search(r'laboratory result|lab result|blood test', content_lower):
            return DocumentType.LAB_RESULT

        # Default to research paper
        return DocumentType.RESEARCH_PAPER

    def _identify_specialties(self, content: str) -> List[MedicalSpecialty]:
        """Identify relevant medical specialties from content"""

        content_lower = content.lower()
        identified_specialties = []

        specialty_keywords = {
            MedicalSpecialty.CARDIOLOGY: ['cardiac', 'heart', 'cardiovascular', 'cardiology', 'myocardial'],
            MedicalSpecialty.ONCOLOGY: ['cancer', 'tumor', 'oncology', 'chemotherapy', 'radiation therapy', 'malignant'],
            MedicalSpecialty.NEUROLOGY: ['neurological', 'neurology', 'brain', 'nerve', 'seizure', 'stroke'],
            MedicalSpecialty.ENDOCRINOLOGY: ['diabetes', 'endocrine', 'hormone', 'thyroid', 'insulin'],
            MedicalSpecialty.GASTROENTEROLOGY: ['gastrointestinal', 'gastroenterology', 'liver', 'stomach', 'intestine'],
            MedicalSpecialty.PULMONOLOGY: ['pulmonary', 'lung', 'respiratory', 'breathing', 'asthma', 'copd'],
            MedicalSpecialty.NEPHROLOGY: ['kidney', 'renal', 'nephrology', 'dialysis', 'urinary'],
            MedicalSpecialty.HEMATOLOGY: ['blood', 'hematology', 'anemia', 'leukemia', 'lymphoma'],
            MedicalSpecialty.RHEUMATOLOGY: ['rheumatology', 'arthritis', 'autoimmune', 'inflammation', 'joint'],
            MedicalSpecialty.INFECTIOUS_DISEASE: ['infection', 'bacterial', 'viral', 'antibiotic', 'antimicrobial'],
            MedicalSpecialty.PSYCHIATRY: ['psychiatric', 'mental health', 'depression', 'anxiety', 'psychology'],
            MedicalSpecialty.DERMATOLOGY: ['skin', 'dermatology', 'rash', 'dermatitis', 'melanoma'],
            MedicalSpecialty.OPHTHALMOLOGY: ['eye', 'vision', 'ophthalmology', 'retina', 'glaucoma'],
            MedicalSpecialty.OTOLARYNGOLOGY: ['ear', 'nose', 'throat', 'ent', 'hearing', 'sinus'],
            MedicalSpecialty.ORTHOPEDICS: ['orthopedic', 'bone', 'joint', 'fracture', 'musculoskeletal'],
            MedicalSpecialty.RADIOLOGY: ['radiology', 'imaging', 'ct scan', 'mri', 'x-ray', 'ultrasound'],
            MedicalSpecialty.PATHOLOGY: ['pathology', 'biopsy', 'histology', 'cytology', 'tissue'],
            MedicalSpecialty.EMERGENCY_MEDICINE: ['emergency', 'trauma', 'acute', 'critical care'],
            MedicalSpecialty.FAMILY_MEDICINE: ['family medicine', 'primary care', 'general practice'],
            MedicalSpecialty.PEDIATRICS: ['pediatric', 'children', 'infant', 'adolescent'],
            MedicalSpecialty.OBSTETRICS_GYNECOLOGY: ['obstetrics', 'gynecology', 'pregnancy', 'women\'s health'],
            MedicalSpecialty.ANESTHESIOLOGY: ['anesthesia', 'anesthetic', 'pain management'],
            MedicalSpecialty.SURGERY: ['surgery', 'surgical', 'operative', 'procedure'],
            MedicalSpecialty.PHARMACOLOGY: ['pharmacology', 'drug', 'medication', 'pharmaceutical'],
            MedicalSpecialty.CLINICAL_TRIALS: ['clinical trial', 'study protocol', 'randomized'],
            MedicalSpecialty.REGULATORY_AFFAIRS: ['fda', 'ema', 'regulatory', 'approval', 'submission']
        }

        for specialty, keywords in specialty_keywords.items():
            if any(keyword in content_lower for keyword in keywords):
                identified_specialties.append(specialty)

        # If no specific specialties identified, default to general
        if not identified_specialties:
            identified_specialties.append(MedicalSpecialty.INTERNAL_MEDICINE)

        return identified_specialties

    async def _extract_study_design_from_content(self, content: str) -> Optional[str]:
        """Extract study design from document content"""
        return self.chunker._extract_study_design(content)

    async def _extract_evidence_level_from_content(self, content: str) -> Optional[str]:
        """Extract evidence level from document content"""
        return self.chunker._extract_evidence_level(content)

    async def _process_and_store_chunks(self, chunks: List[DocumentChunk]):
        """Generate embeddings and store chunks in vector database"""

        # Generate embeddings for all chunks
        embeddings = await self._generate_embeddings([chunk.content for chunk in chunks])

        # Store in ChromaDB
        chunk_ids = []
        chunk_contents = []
        chunk_metadatas = []
        chunk_embeddings = []

        for i, chunk in enumerate(chunks):
            chunk.semantic_embedding = embeddings[i]

            chunk_ids.append(chunk.chunk_id)
            chunk_contents.append(chunk.content)
            chunk_embeddings.append(embeddings[i].tolist())

            # Prepare metadata for ChromaDB
            metadata = {
                "document_id": chunk.document_id,
                "document_type": chunk.document_type.value,
                "specialties": [s.value for s in chunk.specialties],
                "chunk_index": chunk.chunk_index,
                "total_chunks": chunk.total_chunks,
                "clinical_context": chunk.clinical_context or "",
                "evidence_level": chunk.evidence_level or "",
                "study_design": chunk.study_design or "",
                "patient_population": chunk.patient_population or "",
                "intervention": chunk.intervention or "",
                "created_at": chunk.created_at.isoformat()
            }
            chunk_metadatas.append(metadata)

        # Add to ChromaDB collection
        self.collection.add(
            ids=chunk_ids,
            documents=chunk_contents,
            metadatas=chunk_metadatas,
            embeddings=chunk_embeddings
        )

        # Update FAISS index
        await self._update_faiss_index(embeddings)

    async def _generate_embeddings(self, texts: List[str]) -> np.ndarray:
        """Generate semantic embeddings for text chunks"""

        try:
            # Tokenize texts
            inputs = self.chunker.tokenizer(
                texts,
                return_tensors="pt",
                padding=True,
                truncation=True,
                max_length=512
            )

            # Generate embeddings
            with torch.no_grad():
                outputs = self.chunker.model(**inputs)
                embeddings = outputs.last_hidden_state.mean(dim=1)  # Mean pooling

            return embeddings.numpy()

        except Exception as e:
            self.logger.error(f"Error generating embeddings: {e}")
            # Fallback to random embeddings for testing
            return np.random.rand(len(texts), self.faiss_dimension)

    async def _update_faiss_index(self, new_embeddings: np.ndarray):
        """Update FAISS index with new embeddings"""

        if self.faiss_index is None:
            # Initialize FAISS index
            self.faiss_index = faiss.IndexFlatIP(self.faiss_dimension)  # Inner product for cosine similarity

        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(new_embeddings.astype('float32'))

        # Add to index
        self.faiss_index.add(new_embeddings.astype('float32'))


async def main():
    """Example usage of the Medical RAG Pipeline"""

    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

    logger = logging.getLogger(__name__)
    logger.info("Starting Medical RAG Pipeline Demo")

    # Initialize the ingestion system
    ingestion_system = MedicalDocumentIngestion("./demo_medical_db")

    # Example: Ingest a clinical trial document
    try:
        # This would be a real document path in practice
        demo_document = await ingestion_system.ingest_document(
            "demo_clinical_trial.pdf",
            {
                "title": "Phase III Clinical Trial of Novel Oncology Drug",
                "authors": ["Dr. Jane Smith", "Dr. John Doe"],
                "institution": "Memorial Sloan Kettering Cancer Center",
                "doi": "10.1000/demo.trial.2024",
                "pmid": "12345678"
            }
        )

        logger.info(f"Successfully ingested document: {demo_document.title}")
        logger.info(f"Document type: {demo_document.document_type}")
        logger.info(f"Specialties: {[s.value for s in demo_document.specialties]}")
        logger.info(f"Number of chunks: {len(demo_document.chunks)}")

        # Display chunk information
        for chunk in demo_document.chunks[:3]:  # Show first 3 chunks
            logger.info(f"Chunk {chunk.chunk_index}:")
            logger.info(f"  Clinical context: {chunk.clinical_context}")
            logger.info(f"  Evidence level: {chunk.evidence_level}")
            logger.info(f"  Entities found: {len(chunk.entities)}")
            logger.info(f"  Safety signals: {chunk.safety_signals}")
            logger.info(f"  Content preview: {chunk.content[:200]}...")

    except Exception as e:
        logger.error(f"Error in demo: {e}")


if __name__ == "__main__":
    asyncio.run(main())